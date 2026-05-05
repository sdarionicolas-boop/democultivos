# modules/generar_reporte.py
from datetime import datetime
from typing import Tuple, Dict, Any
import io

try:
    from docx import Document
    from docx.shared import Inches, Pt, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from .ia_integration import (
        llamar_deepseek,
        generar_analisis_fertilidad,
        generar_analisis_riesgo_hidrico,
        generar_recomendaciones_integradas,
        preparar_resumen_zonas
    )
    IA_OK = True
except ImportError:
    IA_OK = False

def crear_docx_con_ia(output_path, resultados, cultivo, satelite, fecha_inicio, fecha_fin):
    """
    Versión mejorada que incluye análisis generado por IA.
    """
    doc = Document()
    
    # Título
    title = doc.add_heading(f'REPORTE DE AMBIENTACIÓN AGRONÓMICA - {cultivo}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    # ===== SECCIÓN 1: INTRODUCCIÓN (generada por IA) =====
    doc.add_heading('1. INTRODUCCIÓN', level=1)
    # Preparamos datos para IA
    df_resumen, stats = preparar_resumen_zonas(resultados['gdf_completo'], cultivo)
    
    # Podemos pedir a la IA que genere una introducción personalizada
    if IA_OK:
        prompt_intro = f"""
        Redacta una introducción profesional para un informe agronómico de un lote de {cultivo} de {stats['area_total']:.2f} ha.
        Menciona que se utilizaron imágenes satelitales ({satelite}), modelo digital de elevación y análisis de suelo.
        El objetivo es caracterizar la heterogeneidad del lote y orientar prácticas de manejo específicas.
        """
        intro_text = llamar_deepseek(prompt_intro, temperature=0.5, max_retries=1)
        doc.add_paragraph(intro_text if intro_text else "Análisis no disponible.")
    else:
        doc.add_paragraph("Módulo de IA no disponible. Instale groq para análisis automático.")
    
    # ===== SECCIÓN 2: ANÁLISIS DE FERTILIDAD =====
    doc.add_heading('2. ANÁLISIS DE FERTILIDAD', level=1)
    # Tabla de fertilidad (como ya tienes)
    # ...
    # Luego el análisis de IA
    analisis_fert = generar_analisis_fertilidad(df_resumen, stats, cultivo)
    doc.add_heading('2.1 Interpretación', level=2)
    doc.add_paragraph(analisis_fert)
    
    # ===== SECCIÓN 3: RIESGO HÍDRICO Y TOPOGRAFÍA =====
    doc.add_heading('3. RIESGO DE ENCHARCAMIENTO', level=1)
    # Si hay datos topográficos, mostrarlos...
    # ...
    analisis_agua = generar_analisis_riesgo_hidrico(df_resumen, stats, cultivo)
    doc.add_heading('3.1 Análisis de humedad y textura', level=2)
    doc.add_paragraph(analisis_agua)
    
    # ===== SECCIÓN 4: RECOMENDACIONES INTEGRADAS =====
    doc.add_heading('4. RECOMENDACIONES DE MANEJO', level=1)
    recomendaciones = generar_recomendaciones_integradas(df_resumen, stats, cultivo)
    doc.add_paragraph(recomendaciones)
    
    # ===== SECCIÓN 5: METRICAS DEL LOTE =====
    doc.add_heading('5. MÉTRICAS DEL LOTE', level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Métrica'
    hdr_cells[1].text = 'Valor'
    
    row = table.add_row().cells
    row[0].text = 'Área total'
    row[1].text = f"{stats.get('area_total', 0):.2f} ha"
    
    row = table.add_row().cells
    row[0].text = 'Número de zonas'
    row[1].text = str(stats.get('num_zonas', 0))
    
    row = table.add_row().cells
    row[0].text = 'NDVI promedio'
    row[1].text = f"{stats.get('ndvi_promedio', 0):.3f}"
    
    row = table.add_row().cells
    row[0].text = 'Fecha análisis'
    row[1].text = datetime.now().strftime("%d/%m/%Y")
    
    # ===== SECCIÓN 6: DATOS DEL ANÁLISIS =====
    doc.add_heading('6. DATOS DEL ANÁLISIS', level=1)
    doc.add_paragraph(f"Período: {fecha_inicio} - {fecha_fin}")
    doc.add_paragraph(f"Fuente satelital: {satelite}")
    doc.add_paragraph(f"Cultivo: {cultivo}")
    
    # ===== PIE DE PÁGINA =====
    doc.add_paragraph("")
    p = doc.add_paragraph("生成 por Pachamama - Plataforma de Gestión de Riesgos Climáticos")
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    
    doc.save(output_path)
    return doc


def crear_reporte_simple(output_path: str, resultados: Dict[str, Any], cultivo: str) -> Document:
    """
    Crear reporte basic sin dependence de IA.
    """
    doc = Document()
    
    title = doc.add_heading(f'Reporte de Análisis - {cultivo}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    
    doc.add_heading('Resultados', level=1)
    if 'gdf_completo' in resultados:
        gdf = resultados['gdf_completo']
        doc.add_paragraph(f"Zonas analizadas: {len(gdf)}")
        if hasattr(gdf, 'geometry'):
            area_total = sum(gdf.geometry.area) / 10000
            doc.add_paragraph(f"Área total: {area_total:.2f} ha")
    
    doc.save(output_path)
    return doc
